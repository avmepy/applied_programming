#!/usr/bin/env python3
# -*-encoding: utf-8-*-
# author: Valentyn Kofanov


import docx
from docx.shared import RGBColor





def main(filename_input, filename_output, c1=RGBColor(255, 0, 0), c2=RGBColor(45, 50, 0)):
    doc1 = docx.Document(filename_input)
    doc2 = docx.Document()

    for paragraph in list(doc1.paragraphs):
        p = doc2.add_paragraph()
        for i in range(len(paragraph.runs)):
            run = paragraph.runs[i]
            if run.font.color.rgb == c1:
                run.font.color.rgb = c2
                print(1)

            r = p.add_run()
            r.text = run.text
            r.font.color.rgb = run.font.color.rgb

    doc2.save(filename_output)

if __name__ == '__main__':
    main('res.docx', 'test.docx')