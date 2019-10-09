#!/usr/bin/env python3
# -*-encoding: utf-8-*-
# author: Valentyn Kofanov

import re
import docx


def main(filename):
    doc = docx.Document(filename)
    patt = re.compile(r'-?\d*\.\d*')


    par_list = []

    for par in doc.paragraphs:
        par_list.append(patt.sub(change, par.text))

    res_doc = docx.Document()

    for par in par_list:
        res_doc.add_paragraph(text=par)

    res_doc.save('results.docx')


def change(match: re.match):
    num = match.group()

    if num[0] == '.':
        return '0' + num

    if num[-1] == '.':
        return num + '0'

    return num

if __name__ == '__main__':
    main("main.docx")